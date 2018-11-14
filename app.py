import docx
import os
import sys
import zipfile
import os.path
import lxml.etree
#import PyPDF2
#import
#import pdftableextract as PDF

#### rename file extention from .piz to .zip ####
folder = 'C:/Users/Edwin/Documents/Python/test_plans'
for filename in os.listdir(folder):
    infilename = os.path.join(folder, filename)
    if not os.path.isfile(infilename): continue
    oldbase = os.path.splitext(filename)
    newname = infilename.replace('.piz', '.zip')
    output = os.rename(infilename, newname)

#### unzip .zip file contents to same directory. right now it does not create a new folder
#### still need to program to see and get to file when uploaded, currently hard coded
zip_ref = zipfile.ZipFile('C:/Users/Edwin/Documents/Python/test_plans/CM-215089-VSE881159675HF.zip', 'r')
zip_ref.extractall('C:/Users/Edwin/Documents/Python/test_plans')
zip_ref.close()



#### Load Test Plan Document ####
#### Load Word version ####
doc = docx.Document('AI Test Plan GBS_AISIT7_Final.docx')

#prints the number of paragraphs in the document
#print len(doc.paragraphs)

#### Load PDF version ####
#pdfFileObj = open('AI SIT Test Plan Template_031014.pdf', 'rb')
#pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
#print pdfReader.numPages

#from tabula import read_pdf_table
#table_data = read_pdf_table('AI SIT Test Plan Template_031014.pdf')
#pages = ['3']
#cells = [PDF.process_page("AI SIT Test Plan Template_031014.pdf", p) for p in pages]
#PDF.process_page('AI SIT Test Plan Template_031014.pdf', pages)

#prints the Title of the test plan
print doc.paragraphs[0].text


body_element = doc._body._body
print body_element.xml.encode('utf-8')

toc = lxml.etree.parse(body_element.xml.encode('utf-8'))
print toc.xpath('//w:t[text()="TABLE OF CONTENTS"]')[0].text

#this portion is causing errors
#ps = body_element.xpath('<w:t>TABLE OF CONTENTS</w:t>')
#paragraphs = [paragraph(p, None) for p in ps]


print doc.paragraphs[1].text
print len(doc.paragraphs[1].runs)
print doc.paragraphs[1].runs[0].text
print doc.paragraphs[1].runs[1].text

sections = doc.sections
print sections

print len(sections)

section = sections[0]
print section

for section in sections:
    print(docx.section.Section)


for paragraphs in doc.paragraphs:
    if 'Table 2-2' in paragraphs.text:
        print paragraphs.text

#table = doc.tables[6]

#### Key features & subsystems Table data ####
features_table = doc.tables[5]
features_data = []
for i, row in enumerate(features_table.rows):
    text = (cell.text for cell in row.cells)
    row_data = tuple(text)
    features_data.append(row_data)
print features_data

#### Physical Interface Table data ####
physical_table = doc.tables[6]
physical_data = []
for i, row in enumerate(physical_table.rows):
    text = (cell.text for cell in row.cells)
    row_data = tuple(text)
    physical_data.append(row_data)
print physical_data

#### Logical Interface Table data ####
logical_table = doc.tables[7]
logical_data = []
for i, row in enumerate(logical_table.rows):
    text = (cell.text for cell in row.cells)
    row_data = tuple(text)
    logical_data.append(row_data)
print logical_data
